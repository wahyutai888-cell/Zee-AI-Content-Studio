import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
from PIL import Image

# --- 1. KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Zee AI Studio", layout="wide", page_icon="✨")

# --- 2. MANAJEMEN MULTIPLE API KEYS (ANTI-LIMIT) ---
# Mengambil daftar kunci dari Secrets Streamlit Cloud
# Pastikan di Secrets namanya: GEMINI_KEY_1 dan GEMINI_KEY_2
keys_list = [
    st.secrets.get("GEMINI_KEY_1"), 
    st.secrets.get("GEMINI_KEY_2")
]

# FUNGSI SAKTI UNTUK GENERATE (DENGAN AUTO-SWITCH & FIX MODEL NAME)
def get_ai_response_safe(prompt_data, is_vision=False):
    for i, key in enumerate(keys_list):
        if not key: continue
        try:
            # Konfigurasi ulang setiap ganti kunci
            genai.configure(api_key=key)
            
            # PAKAI NAMA MODEL YANG SUDAH TERBUKTI TEMBUS DI AKUN ZEE
            model_name = 'gemini-flash-latest'
            model = genai.GenerativeModel(model_name)
            
            if is_vision:
                # prompt_data = [instruksi, img]
                response = model.generate_content(prompt_data)
            else:
                # prompt_data = string teks
                response = model.generate_content(prompt_data)
            
            return response.text
        except Exception as e:
            # Jika error 429 (Limit) dan masih ada kunci cadangan, pindah kunci
            if ("429" in str(e) or "quota" in str(e).lower()) and i < len(keys_list) - 1:
                st.sidebar.warning(f"⚠️ Kunci {i+1} Limit, mencoba cadangan...")
                continue
            else:
                # Jika error 404 atau lainnya, tampilkan pesan jelas
                st.error(f"❌ Error Teknis (Model/API): {e}")
                return None
    return None

# --- 3. DESAIN UI AESTHETIC ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600;700&display=swap');
    .stApp { background: radial-gradient(circle at top right, #1e293b, #0f172a); font-family: 'Poppins', sans-serif; color: #e2e8f0; }
    .glass-card { background: rgba(255, 255, 255, 0.03); backdrop-filter: blur(10px); border-radius: 20px; padding: 25px; border: 1px solid rgba(255, 255, 255, 0.1); border-left: 4px solid #3b82f6; }
    .main-title { background: linear-gradient(to right, #60a5fa, #a78bfa); -webkit-background-clip: text; -webkit-text-fill-color: transparent; font-weight: 700; font-size: 3.5rem; text-align: center; }
    .stButton>button { background: linear-gradient(135deg, #3b82f6 0%, #2563eb 100%); color: white; border-radius: 12px; width: 100%; transition: 0.3s; border: none; padding: 10px; font-weight: 600; }
    .stButton>button:hover { transform: translateY(-3px); box-shadow: 0 8px 25px rgba(37, 99, 235, 0.5); }
    </style>
    """, unsafe_allow_html=True)

# --- 4. SESSION STATE ---
if "hasil_caption" not in st.session_state:
    st.session_state.hasil_caption = ""

# --- 5. TAMPILAN UTAMA ---
st.markdown('<h1 class="main-title">Zee AI Content Studio</h1>', unsafe_allow_html=True)
st.markdown('<p style="text-align: center; color: #94a3b8;">Sistem Multimodal AI untuk Affiliate Marketing 2026</p>', unsafe_allow_html=True)

tab1, tab2 = st.tabs(["✨ Script Generator", "👁️ Vision Analysis"])

# --- MODUL 1: SCRIPT ---
with tab1:
    col_kiri, col_kanan = st.columns([1, 1.2], gap="large")
    with col_kiri:
        st.markdown("### 📝 Input Produk")
        prod_info = st.text_area("Detail/Link Produk:", height=150, key="prod_input")
        gaya = st.selectbox("Gaya Bahasa:", ["Persuasif", "Review Jujur", "Hard Sell"])
        if st.button("🚀 Craft My Caption"):
            if prod_info:
                with st.spinner("AI sedang meracik naskah..."):
                    prompt = f"Buat caption affiliate untuk: {prod_info} dengan gaya {gaya}. Langsung berikan teksnya."
                    hasil = get_ai_response_safe(prompt)
                    if hasil: st.session_state.hasil_caption = hasil
            else: st.warning("Masukkan detail produk!")

    with col_kanan:
        st.markdown("### 💎 Hasil Naskah AI")
        if st.session_state.hasil_caption:
            st.markdown(f'<div class="glass-card">{st.session_state.hasil_caption}</div>', unsafe_allow_html=True)
            doc = Document()
            doc.add_heading('Script Affiliate by Zee', 0)
            doc.add_paragraph(st.session_state.hasil_caption)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button("📥 Download Word", data=buffer, file_name="script_zee.docx")

# --- MODUL 2: VISION ---
with tab2:
    st.markdown("### 👁️ AI Vision Studio")
    uploaded_file = st.file_uploader("Upload Foto Produk", type=['png', 'jpg', 'jpeg'])
    if uploaded_file:
        img = Image.open(uploaded_file)
        st.image(img, width=400)
        instruksi = st.text_input("Instruksi AI:", "Analisis foto ini dan berikan deskripsi marketing.")
        if st.button("🔍 Analisis Visual"):
            with st.spinner("Menganalisis gambar..."):
                hasil_v = get_ai_response_safe([instruksi, img], is_vision=True)
                if hasil_v:
                    st.markdown(f'<div class="glass-card">{hasil_v}</div>', unsafe_allow_html=True)
import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
from PIL import Image

# --- 1. KONFIGURASI HALAMAN (Wajib Paling Atas) ---
st.set_page_config(page_title="Zee AI Studio", layout="wide", page_icon="✨")

# --- 2. MANAJEMEN MULTIPLE API KEYS (ANTI-LIMIT) ---
keys_list = [
    st.secrets.get("GEMINI_KEY_1"), 
    st.secrets.get("GEMINI_KEY_2")
]

# FUNGSI SAKTI UNTUK GENERATE (AUTO-SWITCH)
def get_ai_response_safe(prompt_data, is_vision=False):
    for i, key in enumerate(keys_list):
        if not key: continue
        try:
            genai.configure(api_key=key)
            model = genai.GenerativeModel('gemini-flash-latest')
            
            if is_vision:
                response = model.generate_content(prompt_data)
            else:
                response = model.generate_content(prompt_data)
            
            return response.text
        except Exception as e:
            if ("429" in str(e) or "quota" in str(e).lower()) and i < len(keys_list) - 1:
                st.sidebar.warning(f"⚠️ Kunci {i+1} Limit, pindah ke cadangan...")
                continue
            else:
                st.error(f"❌ Error Teknis: {e}")
                return None
    return None

# --- 3. DESAIN UI AESTHETIC (CUSTOM CSS) ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600;700&display=swap');
    .stApp { background: radial-gradient(circle at top right, #1e293b, #0f172a); font-family: 'Poppins', sans-serif; color: #e2e8f0; }
    .glass-card { background: rgba(255, 255, 255, 0.03); backdrop-filter: blur(10px); border-radius: 20px; padding: 25px; border: 1px solid rgba(255, 255, 255, 0.1); box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.3); border-left: 4px solid #3b82f6; }
    .main-title { background: linear-gradient(to right, #60a5fa, #a78bfa); -webkit-background-clip: text; -webkit-text-fill-color: transparent; font-weight: 700; font-size: 3.5rem; text-align: center; margin-bottom: 5px; letter-spacing: -1px; }
    .stButton>button { background: linear-gradient(135deg, #3b82f6 0%, #2563eb 100%); color: white; border: none; padding: 10px 24px; border-radius: 12px; font-weight: 600; letter-spacing: 0.5px; transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1); box-shadow: 0 4px 15px rgba(37, 99, 235, 0.3); width: 100%; }
    .stButton>button:hover { transform: translateY(-3px) scale(1.02); box-shadow: 0 8px 25px rgba(37, 99, 235, 0.5); background: linear-gradient(135deg, #60a5fa 0%, #3b82f6 100%); }
    .stTextInput>div>div>input, .stTextArea>div>div>textarea { background-color: rgba(15, 23, 42, 0.6) !important; border: 1px solid rgba(255, 255, 255, 0.1) !important; color: white !important; border-radius: 12px !important; }
    .stTextInput>div>div>input:focus, .stTextArea>div>div>textarea:focus { border-color: #3b82f6 !important; box-shadow: 0 0 0 2px rgba(59, 130, 246, 0.3) !important; }
    [data-testid="stSidebar"] { background-color: #0f172a; border-right: 1px solid rgba(255, 255, 255, 0.05); }
    .stTabs [data-baseweb="tab-list"] { gap: 10px; background-color: transparent; }
    .stTabs [data-baseweb="tab"] { height: 50px; background-color: rgba(255, 255, 255, 0.05); border-radius: 12px; color: #94a3b8; padding: 0 30px; }
    .stTabs [aria-selected="true"] { background-color: rgba(59, 130, 246, 0.2) !important; color: #60a5fa !important; border: 1px solid #3b82f6 !important; }
    </style>
    """, unsafe_allow_html=True)

# --- 4. MANAJEMEN MEMORI (SESSION STATE) ---
if "hasil_caption" not in st.session_state:
    st.session_state.hasil_caption = ""

# --- 5. TAMPILAN UTAMA ---
st.markdown('<h1 class="main-title">Zee AI Content Studio</h1>', unsafe_allow_html=True)
st.markdown('<p style="text-align: center; color: #94a3b8; font-size: 1.1rem; margin-bottom: 40px;">Sistem Generator Naskah Affiliate & Visual Analysis Terintegrasi.</p>', unsafe_allow_html=True)

tab1, tab2 = st.tabs(["✨ Script Generator", "👁️ Vision Analysis"])

# --- MODUL 1: SCRIPT AFFILIATE ---
with tab1:
    col_kiri, col_kanan = st.columns([1, 1.2], gap="large")
    
    with col_kiri:
        st.markdown("### 📝 Input Produk")
        prod_info = st.text_area("Detail/Link Produk:", placeholder="Contoh: Sepatu lari Nike Air Max, bahan breathable, diskon 50% di Shopee...", height=150)
        
        # Kolom Layout untuk Pilihan (Dikembalikan seperti semula)
        c1, c2 = st.columns(2)
        with c1:
            gaya = st.selectbox("Gaya Bahasa:", ["Persuasif/Mengajak", "Review Jujur", "Hard Sell", "Storytelling"])
        with c2:
            panjang = st.select_slider("Panjang Teks:", options=["Singkat", "Menengah", "Panjang"])
        
        if st.button("🚀 Craft My Caption"):
            if not prod_info:
                st.warning("Mohon masukkan detail produk terlebih dahulu!")
            else:
                with st.spinner("AI sedang meracik naskah..."):
                    # PROMPT SAKTI (Dikembalikan agar hasilnya bagus)
                    prompt = f"""
                    Tugasmu adalah membuat SATU CAPTION MEDIA SOSIAL (Instagram/TikTok) untuk produk affiliate.

                    DETAIL PRODUK: {prod_info}
                    GAYA BAHASA: {gaya}
                    PANJANG: {panjang}

                    INSTRUKSI KHUSUS (WAJIB DIIKUTI):
                    1. JANGAN memberikan kalimat pembuka atau penutup tambahan.
                    2. JANGAN menggunakan label naskah seperti (Opening), (Hook), atau (Closing).
                    3. LANGSUNG berikan teks caption yang siap copy-paste.
                    4. Gunakan Headline yang "hooking" (menangkap perhatian).
                    5. Gunakan poin-poin dengan emoji untuk menjelaskan keunggulan.
                    6. Berikan 5-7 hashtag relevan di paling bawah.
                    """
                    hasil = get_ai_response_safe(prompt)
                    if hasil: st.session_state.hasil_caption = hasil

    with col_kanan:
        st.markdown("### 💎 Hasil Naskah AI")
        if st.session_state.hasil_caption:
            st.markdown(f'<div class="glass-card">{st.session_state.hasil_caption}</div>', unsafe_allow_html=True)
            st.write("") # Spasi
            
            # Export Word
            doc = Document()
            doc.add_heading('Script Affiliate by Zee AI', 0)
            doc.add_paragraph(f"Produk: {prod_info}")
            doc.add_paragraph(f"Gaya: {gaya} | Panjang: {panjang}")
            doc.add_paragraph("-" * 20)
            doc.add_paragraph(st.session_state.hasil_caption)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button("📥 Download to Word (.docx)", data=buffer, file_name="script_affiliate_zee.docx")
        else:
            st.info("Naskah estetikmu akan muncul di sini setelah AI selesai memproses.")

# --- MODUL 2: AI VISION STUDIO ---
with tab2:
    st.markdown("### 👁️ AI Vision Studio")
    st.write("Unggah foto produk untuk dianalisis oleh AI.") # Subteks dikembalikan
    
    uploaded_file = st.file_uploader("Pilih Foto Produk", type=['png', 'jpg', 'jpeg'])
    if uploaded_file:
        c_img, c_res = st.columns([1, 1.5], gap="medium")
        with c_img:
            img = Image.open(uploaded_file)
            st.image(img, use_container_width=True, caption="Foto Produk Terunggah")
        with c_res:
            instruksi = st.text_input("Instruksi untuk AI:", "Berikan deskripsi produk yang estetik dan list fitur yang terlihat di foto ini.")
            if st.button("🔍 Analisis Visual dengan AI"):
                with st.spinner("Vision AI sedang melihat foto..."):
                    hasil_vision = get_ai_response_safe([instruksi, img], is_vision=True)
                    if hasil_vision:
                        st.success("Analisis Selesai!")
                        st.markdown(f'<div class="glass-card">{hasil_vision}</div>', unsafe_allow_html=True)

# --- 6. SIDEBAR ---
with st.sidebar:
    st.markdown("<h2 style='color: #60a5fa;'>🛡️ Zee Admin Panel</h2>", unsafe_allow_html=True)
    status = "Online 🟢" if keys_list[0] else "Offline 🔴"
    st.info(f"Sistem: {status}")
    
    available_keys = sum(1 for k in keys_list if k)
    st.write(f"Kunci Aktif (Load Balancer): {available_keys} Unit")
    
    st.divider()
    st.write("🔍 **System Diagnostics**")
    if st.button("Cek Daftar Model"):
        try:
            # Pengecekan aman
            genai.configure(api_key=keys_list[0])
            models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            st.code(models)
        except Exception as e:
            st.error(f"Gagal: {e}")
            
    st.divider()
    st.caption("© 2026 Skripsi Zee - Sistem Aplikasi Generator Konten Terintegrasi v4.0 (Ultimate Edition)")
# --- 6. SIDEBAR ---
with st.sidebar:
    st.markdown("<h2 style='color: #60a5fa;'>🛡️ Zee Admin Panel</h2>", unsafe_allow_html=True)
    status = "Online 🟢" if keys_list[0] else "Offline 🔴"
    st.info(f"Sistem: {status}")
    
    available_keys = sum(1 for k in keys_list if k)
    st.write(f"Kunci Aktif: {available_keys}")
    
    st.caption("© 2026 Zee AI Studio v3.2.1 - Fix Model Edition")